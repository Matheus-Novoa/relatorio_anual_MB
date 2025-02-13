import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
import openpyxl
from docx2pdf import convert

# Configurações
IMAGES_DIR = "imagens"
EXCEL_DIR = "extracted_tables"
TEMPLATE_PATH = "template.docx"
OUTPUT_DIR_DOCX = "output_docx"
OUTPUT_DIR_PDF = "output_pdf"
JSON_PATH = "pais_com_mais_de_1_filho.json"

# Dicionário de meses em português
MESES_PT = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

# Carregar dados do JSON
with open(JSON_PATH, 'r', encoding='utf-8') as f:
    pais_com_filhos = json.load(f)

def apply_font_settings(doc):
    """Define fonte Arial 13 para todo o documento"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(13)

def process_images():
    today = datetime.now()
    mes_vigente = MESES_PT[today.month]

    for image_file in os.listdir(IMAGES_DIR):
        if image_file.lower().endswith(('.png', '.jpg', '.jpeg')):
            # Extrair nome base e sufixo
            base_name, ext = os.path.splitext(image_file)
            sufixo = None
            responsavel = base_name

            # Verificar sufixos filho1/filho2
            if base_name.endswith('_filho1'):
                responsavel = base_name.rsplit('_filho1', 1)[0]
                sufixo = 'filho1'
            elif base_name.endswith('_filho2'):
                responsavel = base_name.rsplit('_filho2', 1)[0]
                sufixo = 'filho2'
            
            # Carregar dados do Excel
            excel_path = os.path.join(EXCEL_DIR, f"{base_name}.xlsx")
            if not os.path.exists(excel_path):
                print(f"Arquivo Excel não encontrado: {excel_path}")
                continue
                
            wb = openpyxl.load_workbook(excel_path)
            sheet = wb["dados"]
            
            # Extrair dados da primeira linha (assumindo cabeçalho)
            filho_excel = sheet.cell(row=2, column=1).value     # Coluna "Aluno"
            endereco = sheet.cell(row=2, column=2).value  # Coluna "Endereço"
            bairro = sheet.cell(row=2, column=3).value    # Coluna "Bairro"

            # Determinar valor do FILHO
            filho = filho_excel  # Valor padrão
            if responsavel in pais_com_filhos and sufixo in ('filho1', 'filho2'):
                filho = pais_com_filhos[responsavel].get(sufixo, filho_excel).strip()
            
            # Carregar template
            doc = Document(TEMPLATE_PATH)

            # Aplicar formatação Arial 13
            apply_font_settings(doc)
            
            # Substituir placeholders (exceto TABELA)
            replacements = {
                "RESPONSAVEL": responsavel,
                "ENDERECO": endereco,
                "BAIRRO": bairro,
                "FILHO": filho,
                "DIA": str(today.day),
                "MES": mes_vigente,
                "ANO": str(today.year)
            }
            
            # Substituir texto
            for p in doc.paragraphs:
                for key, value in replacements.items():
                    if key in p.text:
                        p.text = p.text.replace(key, str(value))
            
                # Substituir TABELA pela imagem
                    if "TABELA" in p.text:
                        p.text = ""
                        run = p.add_run()
                        run.add_picture(os.path.join(IMAGES_DIR, image_file), width=Inches(6))
                
            # Garantir formatação após alterações
            apply_font_settings(doc)
            
            # Salvar documento
            temp_docx = os.path.join(OUTPUT_DIR_DOCX, f"{base_name}.docx")
            pdf_path = os.path.join(OUTPUT_DIR_PDF, f"{base_name}.pdf")
            
            doc.save(temp_docx) # Salvar DOCX
            convert(temp_docx, pdf_path) # Converter para PDF
            print(f"PDF gerado: {pdf_path}")

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR_DOCX, exist_ok=True)
    os.makedirs(OUTPUT_DIR_PDF, exist_ok=True)
    process_images()